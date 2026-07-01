import io
import base64
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image as PILImage

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False


def convert_word_to_pdf(input_path: str, output_path: str) -> dict:
    if not WEASYPRINT_AVAILABLE:
        raise RuntimeError(
            "WeasyPrint is not installed. Word to PDF conversion requires WeasyPrint. "
            "On Linux: pip install weasyprint, on Windows: install GTK then pip install weasyprint."
        )
    doc = Document(input_path)

    html_parts = [
        '<!DOCTYPE html><html><head><meta charset="utf-8">',
        "<style>",
        """
        @page { margin: 2.54cm; }
        body { font-family: 'Noto Sans', Arial, sans-serif; font-size: 12pt; line-height: 1.6; color: #000; }
        p { margin: 0 0 4pt 0; }
        table { border-collapse: collapse; width: 100%; margin: 8pt 0; }
        td, th { border: 1px solid #999; padding: 4pt 6pt; vertical-align: top; text-align: left; }
        img { max-width: 100%; height: auto; }
        """,
        "</style></head><body>",
    ]

    body_elements = doc.element.body

    for child in body_elements:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = _find_paragraph(doc, child)
            if para:
                html_parts.append(_paragraph_to_html(para, doc))
        elif tag == "tbl":
            table = _find_table(doc, child)
            if table:
                html_parts.append(_table_to_html(table))

    html_parts.append("</body></html>")
    html_str = "".join(html_parts)

    HTML(string=html_str).write_pdf(output_path)
    return {"success": True, "output": output_path}


def _find_paragraph(doc, element):
    for p in doc.paragraphs:
        if p._element is element:
            return p
    return None


def _find_table(doc, element):
    for t in doc.tables:
        if t._element is element:
            return t
    return None


def _paragraph_to_html(paragraph, doc):
    align_map = {
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    align_class = align_map.get(paragraph.alignment, "left")

    parts = []
    for run in paragraph.runs:
        drawings = run._element.findall(qn("w:drawing"))
        if drawings:
            for drawing in drawings:
                img_html = _extract_image_html(drawing, run)
                if img_html:
                    parts.append(img_html)

        text = run.text
        if text:
            style = ""
            if run.bold:
                style += " font-weight:bold;"
            if run.italic:
                style += " font-style:italic;"
            if run.underline:
                style += " text-decoration:underline;"
            if run.font.size:
                style += f" font-size:{run.font.size.pt}pt;"
            if run.font.color and run.font.color.rgb:
                style += f" color:#{run.font.color.rgb};"
            if run.font.name and run.font.name != "Calibri":
                style += f" font-family:'{run.font.name}';"

            if style:
                parts.append(f'<span style="{style}">{_escape(text)}</span>')
            else:
                parts.append(_escape(text))

    html = "".join(parts)
    if not html.strip():
        return "<p><br/></p>"

    styles = []
    pf = paragraph.paragraph_format
    if pf.space_before:
        styles.append(f"margin-top:{pf.space_before.pt}pt")
    if pf.space_after:
        styles.append(f"margin-bottom:{pf.space_after.pt}pt")
    if pf.first_line_indent:
        styles.append(f"text-indent:{pf.first_line_indent.pt}pt")

    style_str = f' style="{";".join(styles)}"' if styles else ""
    return f'<p class="{align_class}"{style_str}>{html}</p>'


def _extract_image_html(drawing_elem, run):
    try:
        blips = drawing_elem.findall(".//" + qn("a:blip"))
        if not blips:
            blips = drawing_elem.findall(".//" + qn("w:blip"))

        for blip in blips:
            embed = blip.get(qn("r:embed"))
            if not embed:
                embed = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
            if embed:
                part = run.part.related_parts.get(embed)
                if part:
                    img_bytes = part.blob
                    pil = PILImage.open(io.BytesIO(img_bytes))
                    if max(pil.size) > 800:
                        ratio = 800 / max(pil.size)
                        pil = pil.resize(
                            (int(pil.width * ratio), int(pil.height * ratio)),
                            PILImage.LANCZOS,
                        )
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    b64 = base64.b64encode(buf.getvalue()).decode()
                    return f'<img src="data:image/png;base64,{b64}" alt="image"/>'
    except Exception:
        pass
    return None


def _table_to_html(table):
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html = []
            for p in cell.paragraphs:
                cell_html.append(_paragraph_text_only(p))
            cells.append(f"<td>{''.join(cell_html)}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f"<table>{''.join(rows)}</table>"


def _paragraph_text_only(paragraph):
    parts = []
    for run in paragraph.runs:
        parts.append(_escape(run.text))
    text = "".join(parts)
    if not text.strip():
        return "<br/>"
    return f"<p>{text}</p>"


def _escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
